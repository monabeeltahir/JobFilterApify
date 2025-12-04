#!/usr/bin/env python3
"""
Resume Parser Module
Extracts skills, experience, and education from resumes (PDF, DOCX, TXT)
"""

import re
from typing import Dict, List, Set, Optional

# Try importing resume parsing libraries
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("PyPDF2 not available. Install with: pip install PyPDF2")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available. Install with: pip install python-docx")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False


class ResumeParser:
    """Parse resumes and extract relevant information"""

    # Common technical skills across various domains
    SKILLS_DATABASE = {
        # Programming Languages
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'swift',
        'kotlin', 'go', 'rust', 'scala', 'r', 'matlab', 'perl', 'shell', 'bash',

        # Web Technologies
        'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask',
        'spring', 'asp.net', 'jquery', 'bootstrap', 'tailwind', 'webpack', 'npm', 'yarn',

        # Databases
        'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite', 'cassandra',
        'dynamodb', 'elasticsearch', 'neo4j', 'firebase',

        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'gitlab', 'github',
        'terraform', 'ansible', 'ci/cd', 'devops', 'cloud',

        # Data Science & ML
        'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'keras', 'scikit-learn',
        'pandas', 'numpy', 'data analysis', 'data science', 'nlp', 'computer vision',
        'artificial intelligence', 'ai', 'ml', 'neural networks',

        # Engineering Disciplines
        'electrical engineering', 'mechanical engineering', 'civil engineering',
        'chemical engineering', 'computer engineering', 'software engineering',
        'systems engineering', 'industrial engineering', 'aerospace engineering',

        # Engineering Skills
        'autocad', 'solidworks', 'matlab', 'simulink', 'plc', 'scada', 'pcb design',
        'circuit design', 'embedded systems', 'microcontrollers', 'fpga', 'vhdl', 'verilog',
        'cad', 'fem', 'cfd', 'ansys', 'catia',

        # Business & Management
        'project management', 'agile', 'scrum', 'kanban', 'product management',
        'business analysis', 'data analysis', 'excel', 'powerpoint', 'tableau',
        'power bi', 'salesforce', 'sap', 'erp', 'crm',

        # Other Technical
        'linux', 'unix', 'windows', 'macos', 'git', 'api', 'rest', 'graphql',
        'microservices', 'testing', 'debugging', 'troubleshooting',
    }

    # Job role keywords
    JOB_ROLES = {
        'software', 'developer', 'engineer', 'programmer', 'architect', 'analyst',
        'scientist', 'researcher', 'manager', 'director', 'lead', 'senior',
        'junior', 'intern', 'consultant', 'specialist', 'technician', 'designer',
        'administrator', 'coordinator', 'supervisor',
    }

    # Education keywords
    EDUCATION_KEYWORDS = {
        'bachelor', 'master', 'phd', 'doctorate', 'degree', 'university', 'college',
        'bs', 'ba', 'ms', 'ma', 'mba', 'engineering', 'science', 'arts', 'technology',
        'computer science', 'electrical engineering', 'mechanical engineering',
        'business administration', 'mathematics', 'physics', 'chemistry',
    }

    def __init__(self):
        self.resume_text = ""
        self.skills = set()
        self.experience_years = 0
        self.job_titles = []
        self.education = []
        self.suggested_roles = []

    def parse_resume(self, file_path: str) -> Dict:
        """
        Parse resume from file and extract information

        Args:
            file_path: Path to resume file (PDF, DOCX, or TXT)

        Returns:
            Dictionary with parsed resume data
        """
        # Extract text from file
        if file_path.lower().endswith('.pdf'):
            self.resume_text = self._extract_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            self.resume_text = self._extract_from_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            self.resume_text = self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format. Supported: PDF, DOCX, TXT")

        if not self.resume_text:
            raise ValueError("Could not extract text from resume")

        # Extract information
        self.skills = self._extract_skills()
        self.experience_years = self._estimate_experience()
        self.job_titles = self._extract_job_titles()
        self.education = self._extract_education()
        self.suggested_roles = self._suggest_job_roles()

        return {
            'text': self.resume_text,
            'skills': list(self.skills),
            'experience_years': self.experience_years,
            'job_titles': self.job_titles,
            'education': self.education,
            'suggested_roles': self.suggested_roles,
            'word_count': len(self.resume_text.split())
        }

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""

        # Try PyPDF2 first
        if PDF_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                if text.strip():
                    return text
            except Exception as e:
                print(f"PyPDF2 extraction failed: {e}")

        # Fallback to pdfplumber
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() + "\n"
                if text.strip():
                    return text
            except Exception as e:
                print(f"pdfplumber extraction failed: {e}")

        if not text.strip():
            raise ValueError("Could not extract text from PDF. Please ensure PDF libraries are installed.")

        return text

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx not available. Install with: pip install python-docx")

        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text

        return text

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()

    def _extract_skills(self) -> Set[str]:
        """Extract skills from resume text"""
        skills = set()
        text_lower = self.resume_text.lower()

        # Find skills from database
        for skill in self.SKILLS_DATABASE:
            # Check for whole word matches or common variations
            pattern = r'\b' + re.escape(skill.replace('-', r'[-\s]?')) + r'\b'
            if re.search(pattern, text_lower, re.IGNORECASE):
                skills.add(skill)

        return skills

    def _estimate_experience(self) -> int:
        """Estimate years of experience from resume"""
        text_lower = self.resume_text.lower()

        # Look for explicit experience mentions
        exp_patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
            r'experience[:\s]+(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s+in',
        ]

        max_years = 0
        for pattern in exp_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    years = int(match)
                    max_years = max(max_years, years)
                except ValueError:
                    pass

        # If no explicit mention, try to count date ranges
        if max_years == 0:
            # Look for date patterns like "2020-2023", "Jan 2020 - Dec 2023"
            date_ranges = re.findall(r'(?:19|20)\d{2}\s*[-–—]\s*(?:(?:19|20)\d{2}|present|current)', text_lower)
            if date_ranges:
                max_years = min(len(date_ranges) * 2, 20)  # Estimate 2 years per job, cap at 20

        return max_years

    def _extract_job_titles(self) -> List[str]:
        """Extract job titles from resume"""
        titles = []
        text_lower = self.resume_text.lower()

        # Common patterns for job titles
        lines = text_lower.split('\n')
        for line in lines:
            line = line.strip()
            # Check if line contains job role keywords
            if any(role in line for role in self.JOB_ROLES):
                # Remove date patterns and clean up
                clean_line = re.sub(r'(?:19|20)\d{2}', '', line)
                clean_line = re.sub(r'[-–—]', ' ', clean_line)
                clean_line = ' '.join(clean_line.split())

                if 10 < len(clean_line) < 100:  # Reasonable length for a job title
                    titles.append(clean_line)

        return titles[:5]  # Return top 5 most likely titles

    def _extract_education(self) -> List[str]:
        """Extract education information from resume"""
        education = []
        text_lower = self.resume_text.lower()

        lines = text_lower.split('\n')
        for line in lines:
            line = line.strip()
            # Check if line contains education keywords
            if any(keyword in line for keyword in self.EDUCATION_KEYWORDS):
                clean_line = ' '.join(line.split())
                if 10 < len(clean_line) < 200:
                    education.append(clean_line)

        return education[:5]  # Return top 5 education entries

    def _suggest_job_roles(self) -> List[str]:
        """Suggest job roles based on skills and experience"""
        suggestions = []
        skills_lower = {skill.lower() for skill in self.skills}

        # Define role mappings based on skills
        role_mappings = {
            'Software Engineer': {'python', 'java', 'javascript', 'programming', 'software'},
            'Data Scientist': {'python', 'machine learning', 'data analysis', 'tensorflow', 'pandas'},
            'DevOps Engineer': {'docker', 'kubernetes', 'aws', 'azure', 'ci/cd', 'jenkins'},
            'Full Stack Developer': {'react', 'node.js', 'javascript', 'html', 'css', 'mongodb'},
            'Frontend Developer': {'react', 'angular', 'vue', 'javascript', 'html', 'css'},
            'Backend Developer': {'python', 'java', 'node.js', 'sql', 'api', 'rest'},
            'Machine Learning Engineer': {'machine learning', 'tensorflow', 'pytorch', 'python', 'deep learning'},
            'Electrical Engineer': {'electrical engineering', 'circuit design', 'pcb', 'embedded systems'},
            'Mechanical Engineer': {'mechanical engineering', 'autocad', 'solidworks', 'cad'},
            'Business Analyst': {'business analysis', 'sql', 'excel', 'data analysis', 'tableau'},
            'Product Manager': {'product management', 'agile', 'scrum', 'project management'},
            'Cloud Engineer': {'aws', 'azure', 'gcp', 'cloud', 'terraform'},
        }

        # Calculate match scores
        role_scores = {}
        for role, required_skills in role_mappings.items():
            match_count = len(required_skills.intersection(skills_lower))
            if match_count > 0:
                score = match_count / len(required_skills)
                role_scores[role] = score

        # Sort by score and return top suggestions
        sorted_roles = sorted(role_scores.items(), key=lambda x: x[1], reverse=True)
        suggestions = [role for role, score in sorted_roles[:5] if score > 0.2]

        return suggestions

    def calculate_job_match_score(self, job_description: str, job_title: str) -> float:
        """
        Calculate how well the resume matches a job posting

        Args:
            job_description: Job description text
            job_title: Job title

        Returns:
            Match score between 0 and 1
        """
        if not self.resume_text:
            return 0.0

        score = 0.0
        job_text_lower = (job_description + " " + job_title).lower()

        # Skill matching (50% weight)
        job_skills = set()
        for skill in self.SKILLS_DATABASE:
            if skill in job_text_lower:
                job_skills.add(skill)

        if job_skills:
            matching_skills = self.skills.intersection(job_skills)
            skill_score = len(matching_skills) / len(job_skills)
            score += skill_score * 0.5

        # Keyword matching (30% weight)
        resume_words = set(self.resume_text.lower().split())
        job_words = set(job_text_lower.split())
        common_words = resume_words.intersection(job_words)
        keyword_score = len(common_words) / max(len(job_words), 1)
        score += min(keyword_score, 1.0) * 0.3

        # Job title matching (20% weight)
        title_score = 0
        for extracted_title in self.job_titles:
            if any(word in job_title.lower() for word in extracted_title.split()):
                title_score = 0.8
                break
        score += title_score * 0.2

        return min(score, 1.0)

    def get_matching_skills(self, job_description: str) -> List[str]:
        """Get skills from resume that match the job description"""
        job_text_lower = job_description.lower()
        matching = []

        for skill in self.skills:
            if skill.lower() in job_text_lower:
                matching.append(skill)

        return matching


def test_parser():
    """Test the resume parser"""
    parser = ResumeParser()
    print("Resume Parser Test")
    print("=" * 80)
    print(f"PDF Support: {PDF_AVAILABLE}")
    print(f"DOCX Support: {DOCX_AVAILABLE}")
    print(f"PDFPlumber Support: {PDFPLUMBER_AVAILABLE}")
    print(f"\nSkills Database Size: {len(parser.SKILLS_DATABASE)}")
    print(f"Job Roles Database Size: {len(parser.JOB_ROLES)}")


if __name__ == "__main__":
    test_parser()
